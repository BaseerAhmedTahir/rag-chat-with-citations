"""Document parsers that preserve page provenance.

Every parser returns a ``ParsedDocument`` whose pages carry 1-indexed page
numbers (PDF) or section indices (DOCX). Page provenance must survive all
the way into the vector store, so it is attached here, at the earliest
possible point.
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

SUPPORTED_EXTENSIONS = {".pdf", ".docx"}


@dataclass(frozen=True)
class PageSpan:
    text: str
    page_number: int  # 1-indexed


@dataclass(frozen=True)
class ParsedDocument:
    source_file: str  # basename of the original file
    pages: list[PageSpan]


def parse_pdf(path: Path) -> ParsedDocument:
    import fitz  # PyMuPDF

    pages: list[PageSpan] = []
    with fitz.open(path) as doc:
        for index, page in enumerate(doc, start=1):
            text = page.get_text("text").strip()
            if text:
                pages.append(PageSpan(text=text, page_number=index))
    return ParsedDocument(source_file=path.name, pages=pages)


def parse_docx(path: Path) -> ParsedDocument:
    """DOCX has no fixed pages; headings delimit sections used as the
    citation unit (section index stands in for page_number)."""
    import docx

    document = docx.Document(str(path))
    sections: list[list[str]] = [[]]
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if paragraph.style.name.startswith("Heading") and sections[-1]:
            sections.append([])
        if text:
            sections[-1].append(text)

    pages = [
        PageSpan(text="\n".join(lines), page_number=index)
        for index, lines in enumerate(sections, start=1)
        if lines
    ]
    return ParsedDocument(source_file=path.name, pages=pages)


def parse_document(path: str | Path) -> ParsedDocument:
    path = Path(path)
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return parse_pdf(path)
    if suffix == ".docx":
        return parse_docx(path)
    raise ValueError(
        f"Unsupported file type {suffix!r} for {path.name}; "
        f"supported: {sorted(SUPPORTED_EXTENSIONS)}"
    )
